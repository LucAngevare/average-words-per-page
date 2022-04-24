import json, re, os, subprocess as subp, time, csv
from datetime import datetime
from docx import Document
from PyPDF2 import PdfFileReader

dictionary = json.load(open("index.json"))

def count_chars(doc):
    newparatextlist = []
    for paratext in doc.paragraphs:
        newparatextlist.append(paratext.text)
    
    return len(re.findall(r'\w+', '\n'.join(newparatextlist)))
dictionary_iteration=0
while dictionary_iteration<len(dictionary):
    for word_amount in range(300, 700):
        doc = Document()
        current_iteration_data = []
        startingTime = datetime.now()
        text = " ".join(dictionary[dictionary_iteration:dictionary_iteration+word_amount])
        paragraph = doc.add_paragraph(text)
        doc.save("sim.docx")
        subp.run('pandoc -o sim.pdf sim.docx', shell=True)
        with open('sim.pdf', 'rb') as f:
            pdf = PdfFileReader(f)
            pageNum = pdf.getNumPages()
        current_iteration_data.extend([dictionary_iteration, count_chars(doc) if count_chars(doc)==word_amount else (count_chars(doc)+word_amount)/2, pageNum, datetime.now().microsecond-startingTime.microsecond if datetime.now().microsecond-startingTime.microsecond>0 else (datetime.now().microsecond-startingTime.microsecond)*-1, len(text)])
        with open('data.csv', 'a', encoding='UTF8', newline='') as f:
            writer = csv.writer(f)
            writer.writerow(current_iteration_data)
        print(str(dictionary_iteration) + ' with a length of ' + str(word_amount) + ' has ' + str(pageNum) + ' pages, and ' + str(len(text)) + ' characters.')
        result = False
        os.remove("sim.docx")
        while not result:
            try:
                os.remove("sim.pdf")
                result = True
            except:
                time.sleep(1)
        if (pageNum == 2):
            dictionary_iteration += word_amount
            break
    else:
        continue